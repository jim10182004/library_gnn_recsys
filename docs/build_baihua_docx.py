"""
把 白話版技術說明.md 轉成 Word docx。

特色：
- 完整支援標題、段落、bullet、numbered list、表格、code block、blockquote、---
- 全程使用 Microsoft JhengHei 中文字型
- 自動處理 **bold**、`inline code`、emoji
- A4 直式、1.5 倍行高、黑體標題彩色

執行：python docs/build_baihua_docx.py
輸出：docs/白話版技術說明.docx
"""
from __future__ import annotations
from pathlib import Path
import re
from docx import Document
from docx.shared import Pt, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT, WD_ALIGN_VERTICAL
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

PROJECT = Path(__file__).parent.parent
SRC = Path(__file__).parent / "白話版技術說明.md"
OUT = Path(__file__).parent / "白話版技術說明.docx"

PRIMARY = RGBColor(0x02, 0x80, 0x90)
DARK = RGBColor(0x1A, 0x2B, 0x33)
GRAY = RGBColor(0x55, 0x55, 0x55)
CODE_BG = RGBColor(0xF4, 0xF4, 0xF4)
CODE_TEXT = RGBColor(0x33, 0x33, 0x33)

CHINESE_FONT = "Microsoft JhengHei"
CODE_FONT = "Consolas"


def set_run_font(run, font_name=CHINESE_FONT, size=11, bold=False, italic=False, color=None):
    run.font.name = font_name
    run.font.size = Pt(size)
    run.font.bold = bold
    run.font.italic = italic
    if color is not None:
        run.font.color.rgb = color
    rpr = run._element.get_or_add_rPr()
    # 移除舊的 rFonts
    for old in rpr.findall(qn("w:rFonts")):
        rpr.remove(old)
    rfonts = OxmlElement("w:rFonts")
    rfonts.set(qn("w:eastAsia"), font_name)
    rfonts.set(qn("w:ascii"), font_name)
    rfonts.set(qn("w:hAnsi"), font_name)
    rpr.insert(0, rfonts)


def set_default_styles(doc):
    style = doc.styles["Normal"]
    style.font.name = CHINESE_FONT
    style.font.size = Pt(11)
    rpr = style.element.get_or_add_rPr()
    rfonts = OxmlElement("w:rFonts")
    rfonts.set(qn("w:eastAsia"), CHINESE_FONT)
    rfonts.set(qn("w:ascii"), CHINESE_FONT)
    rfonts.set(qn("w:hAnsi"), CHINESE_FONT)
    for old in rpr.findall(qn("w:rFonts")):
        rpr.remove(old)
    rpr.insert(0, rfonts)

    sizes = {1: 22, 2: 18, 3: 14, 4: 12}
    for lv in (1, 2, 3, 4):
        try:
            hs = doc.styles[f"Heading {lv}"]
            hs.font.name = CHINESE_FONT
            hs.font.size = Pt(sizes.get(lv, 12))
            hs.font.bold = True
            hs.font.color.rgb = PRIMARY if lv <= 2 else DARK
            rpr = hs.element.get_or_add_rPr()
            for old in rpr.findall(qn("w:rFonts")):
                rpr.remove(old)
            rfonts = OxmlElement("w:rFonts")
            rfonts.set(qn("w:eastAsia"), CHINESE_FONT)
            rfonts.set(qn("w:ascii"), CHINESE_FONT)
            rpr.insert(0, rfonts)
        except KeyError:
            pass


# 內聯文字解析：把一段含 **bold**、`code` 的字串拆成 (text, kind) tuples
TOKEN_RE = re.compile(
    r"(\*\*[^\*\n]+?\*\*)"   # **bold**
    r"|(`[^`\n]+?`)"          # `code`
)


def parse_inline(text):
    """回傳 list of (segment, kind)，kind ∈ {"normal", "bold", "code"}"""
    parts = []
    pos = 0
    for m in TOKEN_RE.finditer(text):
        if m.start() > pos:
            parts.append((text[pos:m.start()], "normal"))
        token = m.group(0)
        if token.startswith("**"):
            parts.append((token[2:-2], "bold"))
        elif token.startswith("`"):
            parts.append((token[1:-1], "code"))
        pos = m.end()
    if pos < len(text):
        parts.append((text[pos:], "normal"))
    if not parts:
        parts = [(text, "normal")]
    return parts


def add_inline_runs(paragraph, text, *, base_size=11, base_color=None, base_bold=False):
    """把 text 的內聯標記轉成 docx runs"""
    for seg, kind in parse_inline(text):
        if not seg:
            continue
        run = paragraph.add_run(seg)
        if kind == "code":
            set_run_font(run, font_name=CODE_FONT, size=base_size - 1,
                         color=CODE_TEXT)
            # 灰底（用 highlight color 較接近）
            rpr = run._element.get_or_add_rPr()
            shd = OxmlElement("w:shd")
            shd.set(qn("w:val"), "clear")
            shd.set(qn("w:color"), "auto")
            shd.set(qn("w:fill"), "F4F4F4")
            rpr.append(shd)
        else:
            set_run_font(
                run, size=base_size,
                bold=base_bold or (kind == "bold"),
                color=base_color,
            )


def add_paragraph(doc, text, *, size=11, color=None, indent_first=False,
                  align=None, space_after=4):
    p = doc.add_paragraph()
    p.paragraph_format.line_spacing = 1.5
    p.paragraph_format.space_after = Pt(space_after)
    if indent_first:
        p.paragraph_format.first_line_indent = Cm(0.74)
    if align is not None:
        p.alignment = align
    add_inline_runs(p, text, base_size=size, base_color=color)
    return p


def add_bullet(doc, text, level=0, size=11):
    p = doc.add_paragraph(style="List Bullet")
    p.paragraph_format.left_indent = Cm(0.74 * (level + 1))
    p.paragraph_format.line_spacing = 1.4
    p.paragraph_format.space_after = Pt(2)
    add_inline_runs(p, text, base_size=size)
    return p


def add_numbered(doc, text, level=0, size=11):
    p = doc.add_paragraph(style="List Number")
    p.paragraph_format.left_indent = Cm(0.74 * (level + 1))
    p.paragraph_format.line_spacing = 1.4
    p.paragraph_format.space_after = Pt(2)
    add_inline_runs(p, text, base_size=size)
    return p


def add_blockquote(doc, text, *, size=11):
    p = doc.add_paragraph()
    p.paragraph_format.left_indent = Cm(1.0)
    p.paragraph_format.line_spacing = 1.4
    p.paragraph_format.space_after = Pt(4)
    pf = p.paragraph_format
    # 淺色背景條的概念用左側邊框模擬
    pPr = p._p.get_or_add_pPr()
    pBdr = OxmlElement("w:pBdr")
    left = OxmlElement("w:left")
    left.set(qn("w:val"), "single")
    left.set(qn("w:sz"), "18")
    left.set(qn("w:space"), "8")
    left.set(qn("w:color"), "028090")
    pBdr.append(left)
    pPr.append(pBdr)
    add_inline_runs(p, text, base_size=size, base_color=GRAY)


def add_code_block(doc, lines, *, size=10):
    """整個 code block 用單一段落 + Consolas + 灰底"""
    p = doc.add_paragraph()
    p.paragraph_format.left_indent = Cm(0.5)
    p.paragraph_format.line_spacing = 1.15
    p.paragraph_format.space_after = Pt(8)
    p.paragraph_format.space_before = Pt(4)
    # 灰色背景（段落級）
    pPr = p._p.get_or_add_pPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:val"), "clear")
    shd.set(qn("w:color"), "auto")
    shd.set(qn("w:fill"), "F4F4F4")
    pPr.append(shd)
    # 加邊框
    pBdr = OxmlElement("w:pBdr")
    for side in ("top", "left", "bottom", "right"):
        b = OxmlElement(f"w:{side}")
        b.set(qn("w:val"), "single")
        b.set(qn("w:sz"), "4")
        b.set(qn("w:space"), "4")
        b.set(qn("w:color"), "DDDDDD")
        pBdr.append(b)
    pPr.append(pBdr)

    text = "\n".join(lines)
    run = p.add_run(text)
    set_run_font(run, font_name=CODE_FONT, size=size, color=CODE_TEXT)


def add_horizontal_rule(doc):
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(6)
    p.paragraph_format.space_before = Pt(6)
    pPr = p._p.get_or_add_pPr()
    pBdr = OxmlElement("w:pBdr")
    bottom = OxmlElement("w:bottom")
    bottom.set(qn("w:val"), "single")
    bottom.set(qn("w:sz"), "8")
    bottom.set(qn("w:space"), "1")
    bottom.set(qn("w:color"), "CCCCCC")
    pBdr.append(bottom)
    pPr.append(pBdr)


def add_heading(doc, text, level):
    if level == 1:
        doc.add_page_break()
    h = doc.add_heading(level=level)
    h.paragraph_format.space_before = Pt(14 if level <= 2 else 8)
    h.paragraph_format.space_after = Pt(8 if level <= 2 else 4)
    add_inline_runs(h, text, base_size={1: 22, 2: 18, 3: 14, 4: 12}.get(level, 11),
                    base_color=PRIMARY if level <= 2 else DARK,
                    base_bold=True)


def parse_table(table_lines):
    """把連續的 table lines 解析成 (headers, rows)"""
    rows = []
    for line in table_lines:
        line = line.strip()
        if not line.startswith("|"):
            continue
        cells = [c.strip() for c in line.strip("|").split("|")]
        rows.append(cells)
    if len(rows) < 2:
        return None, None
    headers = rows[0]
    # rows[1] 是 |---|---|---| 分隔線
    body = rows[2:] if len(rows) > 2 else []
    return headers, body


def add_table(doc, headers, body):
    if not headers:
        return
    n_cols = len(headers)
    table = doc.add_table(rows=1 + len(body), cols=n_cols)
    table.style = "Light Grid Accent 1"
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.autofit = True

    # header
    for j, h in enumerate(headers):
        cell = table.rows[0].cells[j]
        cell.text = ""
        p = cell.paragraphs[0]
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        add_inline_runs(p, h, base_size=10, base_color=DARK, base_bold=True)
        cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER

    # body
    for i, row in enumerate(body):
        for j in range(n_cols):
            val = row[j] if j < len(row) else ""
            cell = table.rows[i + 1].cells[j]
            cell.text = ""
            p = cell.paragraphs[0]
            p.alignment = (WD_ALIGN_PARAGRAPH.LEFT if j == 0
                           else WD_ALIGN_PARAGRAPH.CENTER)
            add_inline_runs(p, val, base_size=10)
            cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER


# ========== 主解析迴圈 ==========

NUMBER_LIST_RE = re.compile(r"^(\d+)\.\s+(.+)$")
BULLET_RE = re.compile(r"^(\s*)[-*+]\s+(.+)$")


def parse_markdown(md_text):
    """逐行解析，回傳 (kind, payload) 列表"""
    blocks = []
    lines = md_text.splitlines()
    i = 0
    n = len(lines)

    while i < n:
        line = lines[i]
        stripped = line.strip()

        # 空行
        if not stripped:
            blocks.append(("blank", None))
            i += 1
            continue

        # Code block ```
        if stripped.startswith("```"):
            i += 1
            code_lines = []
            while i < n and not lines[i].strip().startswith("```"):
                code_lines.append(lines[i])
                i += 1
            blocks.append(("code", code_lines))
            i += 1  # skip closing ```
            continue

        # 標題
        m_h = re.match(r"^(#{1,4})\s+(.+)$", stripped)
        if m_h:
            level = len(m_h.group(1))
            blocks.append(("heading", (level, m_h.group(2).strip())))
            i += 1
            continue

        # 水平線
        if re.match(r"^-{3,}$|^={3,}$|^\*{3,}$", stripped):
            blocks.append(("hr", None))
            i += 1
            continue

        # 表格（連續 | 開頭的行）
        if stripped.startswith("|"):
            tbl_lines = []
            while i < n and lines[i].strip().startswith("|"):
                tbl_lines.append(lines[i])
                i += 1
            blocks.append(("table", tbl_lines))
            continue

        # blockquote
        if stripped.startswith(">"):
            qt = stripped[1:].strip()
            blocks.append(("quote", qt))
            i += 1
            continue

        # numbered list
        m_num = NUMBER_LIST_RE.match(stripped)
        if m_num:
            blocks.append(("numbered", m_num.group(2).strip()))
            i += 1
            continue

        # bullet
        m_b = BULLET_RE.match(line)
        if m_b:
            indent = len(m_b.group(1)) // 2
            blocks.append(("bullet", (indent, m_b.group(2).strip())))
            i += 1
            continue

        # 一般段落
        blocks.append(("para", stripped))
        i += 1

    return blocks


def render(doc, blocks):
    for kind, payload in blocks:
        if kind == "blank":
            continue
        elif kind == "heading":
            level, txt = payload
            add_heading(doc, txt, level)
        elif kind == "hr":
            add_horizontal_rule(doc)
        elif kind == "code":
            add_code_block(doc, payload)
        elif kind == "table":
            headers, body = parse_table(payload)
            if headers:
                add_table(doc, headers, body)
        elif kind == "quote":
            add_blockquote(doc, payload)
        elif kind == "numbered":
            add_numbered(doc, payload)
        elif kind == "bullet":
            level, txt = payload
            add_bullet(doc, txt, level=level)
        elif kind == "para":
            add_paragraph(doc, payload)


def main():
    if not SRC.exists():
        raise FileNotFoundError(f"找不到原稿：{SRC}")

    md = SRC.read_text(encoding="utf-8")
    blocks = parse_markdown(md)

    doc = Document()

    # A4 直式
    section = doc.sections[0]
    section.page_height = Cm(29.7)
    section.page_width = Cm(21.0)
    section.left_margin = Cm(2.2)
    section.right_margin = Cm(2.2)
    section.top_margin = Cm(2.2)
    section.bottom_margin = Cm(2.2)

    set_default_styles(doc)

    # 封面 / 簡頁標題（第一個 H1 會被 page break 推到第 2 頁）
    cover = doc.add_paragraph()
    cover.alignment = WD_ALIGN_PARAGRAPH.CENTER
    for _ in range(6):
        doc.add_paragraph()
    title_p = doc.add_paragraph()
    title_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    title_run = title_p.add_run("我做了什麼？")
    set_run_font(title_run, size=32, bold=True, color=PRIMARY)
    sub_p = doc.add_paragraph()
    sub_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    sub_run = sub_p.add_run("白話版技術說明")
    set_run_font(sub_run, size=22, bold=False, color=DARK)
    for _ in range(2):
        doc.add_paragraph()
    desc_p = doc.add_paragraph()
    desc_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    desc_run = desc_p.add_run("這份文件用日常生活的比喻，解釋我畢業專題的技術細節")
    set_run_font(desc_run, size=12, italic=True, color=GRAY)
    desc_p2 = doc.add_paragraph()
    desc_p2.alignment = WD_ALIGN_PARAGRAPH.CENTER
    desc_run2 = desc_p2.add_run("不需要程式背景也能看懂")
    set_run_font(desc_run2, size=12, italic=True, color=GRAY)
    for _ in range(8):
        doc.add_paragraph()
    foot_p = doc.add_paragraph()
    foot_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    foot_run = foot_p.add_run("基於圖神經網路之市立圖書館書籍推薦系統")
    set_run_font(foot_run, size=12, color=DARK)

    # 渲染主內容
    render(doc, blocks)

    OUT.parent.mkdir(parents=True, exist_ok=True)
    doc.save(str(OUT))
    print(f"已產生：{OUT}")


if __name__ == "__main__":
    main()
